# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from re import finditer

doc = Document()
p = doc.add_paragraph()
text_str = 'dog cat [Customer_Address] dog dog '
for i, ch in enumerate(text_str):
    run = p.add_run(ch)
    font = run.font
    font.name = u'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
    bold = 0
    if i >=9-1 and i <= 9-1+len('[Customer_Address]'):
        bold = 1
    font.bold = bold
print([m.start() for m in finditer('Customer_Address', text_str)])
doc.save('4-2.docx')