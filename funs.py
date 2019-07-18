import docx
from docx.oxml.ns import qn
from re import finditer
from docx.shared import Pt
from docx.shared import RGBColor
from os import path
from os import environ
import win32com.client
from PyPDF2 import PdfFileMerger

def generate_raw_output(name, data_folder, sub_list, output_path = path.join(environ["HOMEPATH"], "Desktop/output.docx")):

	word = win32com.client.Dispatch('Word.Application')
	word.Visible = False
	output = word.Documents.Add()
	files = [data_folder+'/data/Title_Page.docx']
	for i in range(len(sub_list)):
		files.append(data_folder+'/data/'+sub_list[i]+'.docx')

	for i in range(len(files)):
		output.Application.Selection.Range.InsertFile(files[len(files)-i-1])

	output.Range(output.Content.Start, output.Content.End)

	output.SaveAs(output_path+"/"+name+" Warranty Book.docx")
	output.Close()

def fill_information(input_list, output_path = path.join(environ["HOMEPATH"], "Desktop/output.docx")):
	name = input_list[0]
	address = input_list[1]
	Tile1 = input_list[2]
	Tile2 = input_list[3]
	Tile3 = input_list[4]
	Tile4 = input_list[5]
	Tile5 = input_list[6]
	document = docx.Document(output_path+"/"+name+" Warranty Book.docx")
	for paragraph in document.paragraphs:
		if 'Customer_Address' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Customer_Address]",address)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(address, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				if place != 0:
					run = paragraph.add_run(ch)
					font = run.font
					font.name = u'Calibri (Body)'
					run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri (Body)')
					bold = 0
					if i >= place - 1 and i <= place - 1 + len(address):
						bold = 1
					font.bold = bold
				else:
					run = paragraph.add_run(ch)
					font = run.font
					font.name = u'Arial'
					font.size = Pt(24)
					color = font.color
					color.rgb = RGBColor(255,255,255)
					run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Customer_Name' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Customer_Name]",name)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(name, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(24)
				color = font.color
				color.rgb = RGBColor(255, 255, 255)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Tile_1' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Tile_1]",Tile1)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(Tile1, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(11)
				color = font.color
				color.rgb = RGBColor(0,0,0)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Tile_2' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Tile_2]",Tile2)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(Tile2, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(11)
				color = font.color
				color.rgb = RGBColor(0,0,0)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Tile_3' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Tile_3]",Tile3)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(Tile3, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(11)
				color = font.color
				color.rgb = RGBColor(0,0,0)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Tile_4' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Tile_4]",Tile4)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(Tile4, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(11)
				color = font.color
				color.rgb = RGBColor(0,0,0)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
		if 'Tile_5' in paragraph.text:
			paragraph.text = paragraph.text.replace("[Tile_5]",Tile5)
			paragraph_reserve = paragraph.text
			paragraph.text = ""
			try:
				place = [m.start() for m in finditer(Tile5, paragraph_reserve)][0]
			except:
				place = 0
			for i, ch in enumerate(paragraph_reserve):
				run = paragraph.add_run(ch)
				font = run.font
				font.name = u'Arial'
				font.size = Pt(11)
				color = font.color
				color.rgb = RGBColor(0,0,0)
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Arial')
	document.save(output_path+"/"+name+" Warranty Book.docx")

def pdf_combiner(data_folder, PC_list, name, output_path = path.join(environ["HOMEPATH"], "Desktop/output.docx")): # Combine the files in the pdf
	files = []
	for i in range(len(PC_list)):
		if "Reece" in PC_list[i]:
			files.append(data_folder + '/data/Reece/'+PC_list[i].replace(" Reece", "").upper())
		if "Southern" in PC_list[i]:
			files.append(data_folder + '/data/Southern/'+PC_list[i].replace(" Southern", "").upper())

	merger = PdfFileMerger()
	for i in range(len(files)):
		merger.append(open(files[i], 'rb'))
	with open(output_path + "/" + name + " Warranty Book PC Items.pdf", 'wb') as fout:  # different from func
		# This is the output path
		merger.write(fout)